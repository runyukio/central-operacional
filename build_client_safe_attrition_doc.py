from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/lucaskawakami/Documents/New project")
OUTPUT = ROOT / "outputs" / "attrition_diagnostic_client_safe.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 35, 45)
GRAY = RGBColor(95, 95, 95)
LIGHT_GRAY = "F2F4F7"
CALL_OUT = "F4F6F9"
BORDER = "D9E2EC"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 23, RGBColor(0, 0, 0)),
        ("Subtitle", 14, GRAY),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, color=BLUE, bold=True)
    return paragraph


def add_para(doc, text, bold=False, color=INK, size=11, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [Inches(6.5 / len(metrics))] * len(metrics)
    set_table_width(table, widths)
    set_table_borders(table, color="FFFFFF")
    for idx, (label, value) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, CALL_OUT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        label_run = p.add_run(label.upper() + "\n")
        set_run_font(label_run, size=8.5, color=GRAY, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=15, color=DARK_BLUE, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, widths, number_cols=None):
    number_cols = number_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if idx in number_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10.5, color=INK, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if idx in number_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=10.5, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_note_box(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [Inches(6.4)])
    set_table_borders(table, color="E1E7EF")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALL_OUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title + "\n")
    set_run_font(title_run, size=10.5, color=DARK_BLUE, bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Attrition Diagnostic"
    set_run_font(header.runs[0], size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Prepared May 20, 2026")
    set_run_font(footer_run, size=9, color=GRAY)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("ATTRITION DIAGNOSTIC")
    set_run_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Training and Early Production Lifecycle Review")
    set_run_font(run, size=14, color=GRAY)

    for label, value in [
        ("Scope:", "Aggregated view of one analyzed LOB, with benchmark context from other LOBs"),
        ("Focus:", "Training attrition and production attrition after Go Live"),
        ("Privacy approach:", "No employee names, file names, or internal contract-mix details included"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(label + " ")
        set_run_font(label_run, size=11, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=11, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_metric_strip(
        doc,
        [
            ("Started", "84"),
            ("Training attrition", "21 / 25.0%"),
            ("Production attrition", "21 / 25.0%"),
            ("Total attrition", "42 / 50.0%"),
        ],
    )

    add_heading(doc, "Executive Summary", 1)
    add_bullet(doc, "Attrition is concentrated in two moments: training and the first weeks after Go Live.")
    add_bullet(doc, "The strongest signal is voluntary attrition. In the analyzed LOB, 37 of 42 exits were voluntary.")
    add_bullet(doc, "The data does not support shift as the only root cause. Turnover is similar across Morning, Afternoon, and Night.")
    add_bullet(doc, "Training attrition is concentrated in specific waves, especially Wave 35.")
    add_bullet(doc, "Production attrition happens early. The median time after Go Live is 13 days.")

    add_note_box(
        doc,
        "Client-safe takeaway",
        "The main opportunity is to strengthen expectation alignment before start, improve readiness checks during training, and add closer support in the first two weeks after Go Live.",
    )

    add_heading(doc, "Training Attrition", 1)
    add_para(
        doc,
        "Training attrition represents 21 exits out of 84 employees who started, equal to 25.0%. The issue is not evenly distributed across all waves; it is concentrated in a few cohorts.",
    )
    add_table(
        doc,
        ["Wave", "Started", "Training Exits", "Training Attrition"],
        [
            ["Wave 27", 5, 0, "0.0%"],
            ["Wave 28", 1, 1, "100.0%"],
            ["Wave 29", 8, 0, "0.0%"],
            ["Wave 31", 13, 0, "0.0%"],
            ["Wave 32", 8, 1, "12.5%"],
            ["Wave 33", 18, 7, "38.9%"],
            ["Wave 34", 17, 4, "23.5%"],
            ["Wave 35", 14, 8, "57.1%"],
        ],
        [Inches(1.3), Inches(1.3), Inches(1.7), Inches(1.7)],
        number_cols={1, 2, 3},
    )
    add_para(
        doc,
        "Wave 35 is the clearest training alert. While it was an Afternoon cohort, the stated reasons point more to offer competitiveness, financial viability, personal constraints, and fit than to shift alone.",
    )

    add_heading(doc, "Production Attrition", 1)
    add_para(
        doc,
        "Production attrition also represents 21 exits out of 84 employees who started, equal to 25.0%. The key issue is timing: exits happen very early after Go Live.",
    )
    add_table(
        doc,
        ["Wave", "Post-Go Live Exits", "Production Attrition", "Median Days After Go Live"],
        [
            ["Wave 27", 3, "60.0%", "75 days"],
            ["Wave 29", 3, "37.5%", "10 days"],
            ["Wave 31", 6, "46.2%", "25 days"],
            ["Wave 32", 3, "37.5%", "11 days"],
            ["Wave 33", 1, "5.6%", "55 days"],
            ["Wave 34", 5, "29.4%", "12 days"],
            ["Wave 35", 0, "0.0%", "Not mature yet"],
        ],
        [Inches(1.15), Inches(1.65), Inches(1.6), Inches(2.1)],
        number_cols={1, 2, 3},
    )
    add_table(
        doc,
        ["Time After Go Live", "Production Exits", "% of Production Exits"],
        [
            ["0 days", 2, "9.5%"],
            ["1-7 days", 3, "14.3%"],
            ["8-14 days", 6, "28.6%"],
            ["15-30 days", 2, "9.5%"],
            ["31-60 days", 6, "28.6%"],
            ["61+ days", 2, "9.5%"],
        ],
        [Inches(2.5), Inches(1.7), Inches(1.8)],
        number_cols={1, 2},
    )
    add_para(
        doc,
        "More than half of production exits happened within 14 days after Go Live. This suggests the transition from training to production is a critical retention moment.",
    )

    add_heading(doc, "Shift View", 1)
    add_table(
        doc,
        ["Shift", "Started", "Total Exits", "Total Attrition"],
        [
            ["Morning", 21, 11, "52.4%"],
            ["Afternoon", 24, 12, "50.0%"],
            ["Night", 39, 19, "48.7%"],
        ],
        [Inches(1.6), Inches(1.4), Inches(1.4), Inches(1.6)],
        number_cols={1, 2, 3},
    )
    add_para(
        doc,
        "Shift should be treated as a lens, not as the sole cause. Attrition levels are similar across shifts. The main exception is Wave 35, where the Afternoon cohort concentrated training exits.",
    )

    add_heading(doc, "Reasons Captured", 1)
    add_para(
        doc,
        "The captured reasons are directional and should be used as qualitative evidence, not as a complete reason universe for all exits.",
    )
    add_table(
        doc,
        ["Reason Category", "Cases", "% of Captured Reasons"],
        [
            ["Another job offer", 5, "35.7%"],
            ["Personal or schedule constraints", 4, "28.6%"],
            ["Financial, commute, or training disruption", 3, "21.4%"],
            ["Job fit or expectation mismatch", 2, "14.3%"],
        ],
        [Inches(3.4), Inches(1.2), Inches(1.7)],
        number_cols={1, 2},
    )
    add_para(
        doc,
        "These reasons reinforce that the main opportunity is candidate alignment and early experience, rather than performance management alone.",
    )

    add_heading(doc, "Benchmark Context", 1)
    add_table(
        doc,
        ["Population", "Primary Signal", "What It Suggests"],
        [
            ["Analyzed LOB - current cycle", "Voluntary exits dominate", "Retention, expectation alignment, and early experience"],
            ["Benchmark training exits", "Voluntary exits also dominate", "Training attrition often reflects fit and offer acceptance risk"],
            ["Benchmark production exits", "Involuntary exits dominate", "More related to performance or operational fit"],
        ],
        [Inches(2.2), Inches(1.9), Inches(2.3)],
    )
    add_para(
        doc,
        "This benchmark indicates that the analyzed LOB has a distinct voluntary attrition pattern in the current cycle. It should not be reduced to one factor such as shift or contract model without further validation.",
    )

    add_heading(doc, "Root Cause Hypotheses", 1)
    add_table(
        doc,
        ["Hypothesis", "Evidence", "Risk If Not Addressed"],
        [
            ["Expectation mismatch", "Reasons mention fit and adaptation challenges", "Candidates leave before stabilizing"],
            ["Offer competitiveness", "Several exits mention another job offer", "Accepted candidates remain open to alternatives"],
            ["Financial or commute viability", "Reasons mention financial constraints and transportation/training disruption", "Early voluntary exits continue"],
            ["Early production adaptation", "Median production tenure after Go Live is 13 days", "Losses continue immediately after handoff"],
        ],
        [Inches(1.8), Inches(2.4), Inches(2.0)],
    )

    add_heading(doc, "Recommended Actions", 1)
    add_bullet(doc, "Strengthen pre-start alignment: schedule, pay expectations, work routine, commute, and availability.")
    add_bullet(doc, "Add a readiness check before training starts to confirm the candidate still accepts the full offer conditions.")
    add_bullet(doc, "Review Wave 35 learnings before repeating the same hiring and training setup.")
    add_bullet(doc, "Create structured check-ins during training and during the first two weeks after Go Live.")
    add_bullet(doc, "Track training attrition and production attrition separately going forward.")

    add_note_box(
        doc,
        "How to position this externally",
        "The data points to a lifecycle retention opportunity: improve alignment before start, reinforce support during training, and reduce early production losses after Go Live.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
