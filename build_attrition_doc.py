from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/lucaskawakami/Documents/New project")
INPUT = Path("/Users/lucaskawakami/Downloads/Turnnover.xlsx")
REASONS_INPUT = Path("/Users/lucaskawakami/Downloads/ADS desistências.xlsx")
OUTPUT = ROOT / "outputs" / "attrition_insights_report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
CALL_OUT = "F4F6F9"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_paragraph(doc, text="", style=None, bold=False, color=INK, size=11, after=6):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [Inches(6.5 / len(metrics))] * len(metrics)
    set_table_width(table, widths)
    set_table_borders(table, color="FFFFFF")
    for idx, (label, value) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, CALL_OUT)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        label_run = p.add_run(label.upper() + "\n")
        set_run_font(label_run, size=8.5, color=GRAY, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=15, color=DARK_BLUE, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, widths, number_cols=None):
    number_cols = number_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_GRAY)
        paragraph = hdr[idx].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        if idx in number_cols:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=10.5, color=INK, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            if idx in number_cols:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            set_run_font(run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 23, RGBColor(0, 0, 0)),
        ("Subtitle", 14, GRAY),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def read_post_training_metrics():
    df = pd.read_excel(INPUT, sheet_name="Planilha1")
    for column in ["Data de Contratação", "Data Nesting", "Go live", "Data de Desligamento"]:
        df[column] = pd.to_datetime(df[column], errors="coerce")
    for column in ["Recrutamento Total", "Turnover Pos", "Turnover Treinamento"]:
        df[column] = pd.to_numeric(df[column], errors="coerce").fillna(0).astype(int)

    df["Exited"] = ((df["Turnover Pos"] == 1) | (df["Turnover Treinamento"] == 1)).astype(int)
    df["ActiveFlag"] = df["Ativo"].astype(str).str.lower().eq("active").astype(int)
    df["PostTenureDays"] = (df["Data de Desligamento"] - df["Go live"]).dt.days
    post = df[df["Turnover Pos"] == 1].copy()
    training = df[df["Turnover Treinamento"] == 1].copy()
    total_started = int(df["Recrutamento Total"].sum())
    total_exits = int(df["Exited"].sum())

    wave_summary = df.groupby("Wave").agg(
        started=("Recrutamento Total", "sum"),
        active=("ActiveFlag", "sum"),
        training=("Turnover Treinamento", "sum"),
        post=("Turnover Pos", "sum"),
        total_exit=("Exited", "sum"),
    )
    wave_summary["training_rate"] = wave_summary["training"] / wave_summary["started"]
    wave_summary["post_rate"] = wave_summary["post"] / wave_summary["started"]
    wave_summary["total_rate"] = wave_summary["total_exit"] / wave_summary["started"]
    wave_summary["avg_post_tenure_days"] = post.groupby("Wave")["PostTenureDays"].mean()
    wave_summary["median_post_tenure_days"] = post.groupby("Wave")["PostTenureDays"].median()

    bins = [-9999, 0, 7, 14, 30, 60, 90, 180, 9999]
    labels = ["0 days", "1-7 days", "8-14 days", "15-30 days", "31-60 days", "61-90 days", "91-180 days", "181+ days"]
    post["PostBucket"] = pd.cut(post["PostTenureDays"], bins=bins, labels=labels)
    post_bucket_counts = post["PostBucket"].value_counts().sort_index()

    stage_type_rows = []
    for label, frame in [("Training", training), ("Post-training", post), ("Total exits", df[df["Exited"] == 1])]:
        counts = frame["Type of Resign"].value_counts()
        voluntary = int(counts.get("Voluntary", 0))
        involuntary = int(counts.get("Involuntary", 0))
        total = voluntary + involuntary
        stage_type_rows.append(
            {
                "stage": label,
                "involuntary": involuntary,
                "voluntary": voluntary,
                "total": total,
                "voluntary_rate": voluntary / total if total else 0,
            }
        )
    reasons_metrics = read_reasons_metrics(df)

    return {
        "total_started": total_started,
        "active": int(df["ActiveFlag"].sum()),
        "training_exits": int(training.shape[0]),
        "post_exits": int(post.shape[0]),
        "total_exits": total_exits,
        "training_rate": training.shape[0] / total_started,
        "post_rate": post.shape[0] / total_started,
        "total_rate": total_exits / total_started,
        "min_go_live": df["Go live"].min(),
        "max_go_live": df["Go live"].max(),
        "avg_post_days": post["PostTenureDays"].mean(),
        "median_post_days": post["PostTenureDays"].median(),
        "min_post_days": post["PostTenureDays"].min(),
        "max_post_days": post["PostTenureDays"].max(),
        "wave_summary": wave_summary,
        "post_bucket_counts": post_bucket_counts,
        "stage_type_rows": stage_type_rows,
        "reasons": reasons_metrics,
    }


def pct(value):
    return f"{value * 100:.1f}%"


def normalize_name(value):
    import re
    import unicodedata

    text = str(value).replace("\xa0", " ")
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = re.sub(r"[^A-Z0-9 ]+", " ", text.upper())
    return re.sub(r"\s+", " ", text).strip()


def reason_category(reason):
    text = normalize_name(reason)
    if "OUTRA PROPOSTA" in text:
        return "Another job offer"
    if "FINANCEIR" in text or "BILHETE" in text or "RECURSOS" in text or "SUSPENSO" in text:
        return "Financial / commute / training disruption"
    if "EXPECTATIV" in text or "ADAPTAR" in text:
        return "Job fit or expectation mismatch"
    if "PESSOA" in text or "INVIAVEL" in text or "RELACIONAMENTO" in text or "TURNO" in text:
        return "Personal or schedule constraints"
    return "Other / not specified"


def read_reasons_metrics(turnover_df):
    raw = pd.read_excel(REASONS_INPUT, sheet_name="Planilha1", header=None)
    reasons = raw.iloc[2:, 0:4].copy()
    reasons.columns = ["Nome", "Wave", "Data", "Motivo"]
    reasons = reasons.dropna(subset=["Nome", "Wave", "Data", "Motivo"]).copy()
    reasons["Data"] = pd.to_datetime(reasons["Data"], errors="coerce")
    reasons["Category"] = reasons["Motivo"].apply(reason_category)
    reasons["key"] = reasons["Nome"].apply(normalize_name)

    turnover_match = turnover_df.copy()
    turnover_match["key"] = turnover_match["Nome"].apply(normalize_name)
    merged = reasons.merge(
        turnover_match[
            [
                "key",
                "Nome",
                "Wave",
                "Turnover Pos",
                "Turnover Treinamento",
                "Type of Resign",
                "Data de Desligamento",
            ]
        ],
        on="key",
        how="left",
        suffixes=("_reason", "_turnover"),
    )
    merged["Matched"] = merged["Nome_turnover"].notna()

    category_counts = reasons["Category"].value_counts()
    wave_counts = reasons.groupby("Wave").size()
    matched = merged[merged["Matched"]]
    training_matched = int((matched["Turnover Treinamento"] == 1).sum())
    post_matched = int((matched["Turnover Pos"] == 1).sum())
    unmatched_names = merged.loc[~merged["Matched"], "Nome_reason"].tolist()

    return {
        "total_reasons": int(len(reasons)),
        "category_counts": category_counts,
        "wave_counts": wave_counts,
        "matched_count": int(merged["Matched"].sum()),
        "training_matched": training_matched,
        "post_matched": post_matched,
        "unmatched_count": int((~merged["Matched"]).sum()),
        "unmatched_names": unmatched_names,
    }


def main():
    metrics = read_post_training_metrics()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "Attrition Insights Report"
    header_p.paragraph_format.space_after = Pt(0)
    set_run_font(header_p.runs[0], size=9, color=GRAY)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Prepared May 20, 2026")
    set_run_font(footer_run, size=9, color=GRAY)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("ATTRITION INSIGHTS REPORT")
    set_run_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Training Attrition and Post-Training Implementation Analysis")
    set_run_font(run, size=14, color=GRAY)

    for label, value in [
        ("Source:", "Turnnover.xlsx"),
        ("Scope:", f"{metrics['total_started']} employees who started across the listed waves"),
        ("Period:", f"Go Live dates from {metrics['min_go_live']:%b %d, %Y} to {metrics['max_go_live']:%b %d, %Y}"),
        ("Purpose:", "Summarize where attrition is happening and what to prioritize first"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(label + " ")
        set_run_font(label_run, size=11, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=11, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_metric_strip(
        doc,
        [
            ("Total started", f"{metrics['total_started']}"),
            ("Total attrition", f"{metrics['total_exits']} ({pct(metrics['total_rate'])})"),
            ("Training exits", f"{metrics['training_exits']} ({pct(metrics['training_rate'])})"),
            ("Post-training exits", f"{metrics['post_exits']} ({pct(metrics['post_rate'])})"),
        ],
    )

    add_heading(doc, "Executive Summary", 1)
    add_bullet(doc, f"{metrics['total_started']} employees started across the waves in the file. {metrics['active']} are still active and {metrics['total_exits']} have left.")
    add_bullet(doc, f"Total attrition is {pct(metrics['total_rate'])}. This is split evenly between training attrition and post-training attrition: {metrics['training_exits']} exits in each stage.")
    add_bullet(doc, f"Post-training exits happened quickly after Go Live. The average time after Go Live is {metrics['avg_post_days']:.0f} days and the median is {metrics['median_post_days']:.0f} days.")
    add_bullet(doc, "Most exits are voluntary: 37 out of 42 exits, or 88.1% of total attrition.")
    add_bullet(doc, f"The additional reasons file explains {metrics['reasons']['total_reasons']} exits. The most common reason category is another job offer, followed by personal or schedule constraints.")
    add_bullet(doc, "Training attrition is concentrated in Waves 33, 34, and 35. Together, they represent 19 of the 21 training exits.")
    add_bullet(doc, "Post-training attrition is highest by volume in Wave 31 and Wave 34. Wave 27 has the highest post-training rate, but it has a small base of 5 people.")

    add_heading(doc, "Important Data Note", 1)
    add_paragraph(
        doc,
        "This report uses Recrutamento Total as the number of people who started in each wave. Turnover Treinamento = 1 means the employee left during training. Turnover Pos = 1 means the employee left after training. For post-training exits, time in the company after implementation is calculated as Data de Desligamento minus Go live.",
    )

    add_heading(doc, "Overall Attrition Split", 1)
    overall_rows = [
        ["Started", metrics["total_started"], "100.0%"],
        ["Still active", metrics["active"], pct(metrics["active"] / metrics["total_started"])],
        ["Training exits", metrics["training_exits"], pct(metrics["training_rate"])],
        ["Post-training exits", metrics["post_exits"], pct(metrics["post_rate"])],
        ["Total exits", metrics["total_exits"], pct(metrics["total_rate"])],
    ]
    add_table(
        doc,
        ["Category", "Employees", "% of Started"],
        overall_rows,
        [Inches(2.7), Inches(1.5), Inches(1.6)],
        number_cols={1, 2},
    )

    add_heading(doc, "Turnover by Wave", 1)
    wave_rows = []
    for wave, row in metrics["wave_summary"].sort_index().iterrows():
        wave_rows.append(
            [
                wave,
                int(row["started"]),
                int(row["active"]),
                f"{int(row['training'])} ({pct(row['training_rate'])})",
                f"{int(row['post'])} ({pct(row['post_rate'])})",
                f"{int(row['total_exit'])} ({pct(row['total_rate'])})",
            ]
        )
    add_table(
        doc,
        ["Wave", "Started", "Active", "Training Turnover", "Post Turnover", "Total Turnover"],
        wave_rows,
        [Inches(0.85), Inches(0.85), Inches(0.8), Inches(1.55), Inches(1.35), Inches(1.4)],
        number_cols={1, 2, 3, 4, 5},
    )
    add_paragraph(
        doc,
        "Wave 35 has the highest meaningful training turnover rate among larger waves: 8 of 14 employees, or 57.1%. Wave 28 shows 100.0%, but the base is only one employee. Post-training turnover is highest by rate in Wave 27, with 3 of 5 employees, or 60.0%.",
    )

    add_heading(doc, "Voluntary vs. Involuntary", 1)
    type_rows = [
        [
            row["stage"],
            row["involuntary"],
            row["voluntary"],
            row["total"],
            pct(row["voluntary_rate"]),
        ]
        for row in metrics["stage_type_rows"]
    ]
    add_table(
        doc,
        ["Stage", "Involuntary", "Voluntary", "Total Exits", "% Voluntary"],
        type_rows,
        [Inches(1.65), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.25)],
        number_cols={1, 2, 3, 4},
    )
    add_paragraph(
        doc,
        "The biggest behavioral pattern is voluntary attrition. This suggests the main investigation should focus on expectation setting, schedule and compensation fit, job reality, leadership support, and early employee experience.",
    )

    add_heading(doc, "Reasons Provided", 1)
    reason_rows = []
    for category, count in metrics["reasons"]["category_counts"].items():
        reason_rows.append([category, int(count), pct(count / metrics["reasons"]["total_reasons"])])
    add_table(
        doc,
        ["Reason Category", "Employees", "% of Reasons"],
        reason_rows,
        [Inches(3.8), Inches(1.2), Inches(1.4)],
        number_cols={1, 2},
    )
    wave_reason_rows = []
    for wave, count in metrics["reasons"]["wave_counts"].items():
        wave_reason_rows.append([wave, int(count), pct(count / metrics["reasons"]["total_reasons"])])
    add_table(
        doc,
        ["Wave", "Reasons Captured", "% of Reasons"],
        wave_reason_rows,
        [Inches(1.6), Inches(1.8), Inches(1.6)],
        number_cols={1, 2},
    )
    add_paragraph(
        doc,
        f"The reasons file contains {metrics['reasons']['total_reasons']} records. {metrics['reasons']['matched_count']} matched employees in the turnover base: {metrics['reasons']['training_matched']} training exits and {metrics['reasons']['post_matched']} post-training exits. {metrics['reasons']['unmatched_count']} record did not match the turnover base.",
    )
    if metrics["reasons"]["unmatched_names"]:
        add_paragraph(
            doc,
            "Unmatched reason record: " + ", ".join(metrics["reasons"]["unmatched_names"]) + ".",
        )
    add_paragraph(
        doc,
        "The most frequent specific theme is another job offer. This points to competitiveness risk: candidates may be using training as a bridge while evaluating other options, or the offer may not be strong enough compared with alternatives.",
    )

    add_heading(doc, "Post-Training Time After Go Live", 1)
    post_tenure_rows = []
    for wave, row in metrics["wave_summary"].sort_index().iterrows():
        if int(row["post"]) == 0:
            continue
        avg = row["avg_post_tenure_days"]
        median = row["median_post_tenure_days"]
        post_tenure_rows.append(
            [
                wave,
                int(row["post"]),
                f"{avg:.0f} days",
                f"{median:.0f} days",
            ]
        )
    add_table(
        doc,
        ["Wave", "Post Exits", "Avg. Days After Go Live", "Median Days After Go Live"],
        post_tenure_rows,
        [Inches(1.1), Inches(1.2), Inches(2.1), Inches(2.1)],
        number_cols={1, 2, 3},
    )
    add_paragraph(
        doc,
        f"Across all post-training exits, the average time after Go Live is {metrics['avg_post_days']:.0f} days and the median is {metrics['median_post_days']:.0f} days. The median is much lower than the average because a few longer-tenure exits pull the average upward.",
    )

    add_heading(doc, "Post-Training Exit Timing", 1)
    bucket_rows = []
    for bucket, count in metrics["post_bucket_counts"].items():
        bucket_rows.append(
            [
                bucket,
                int(count),
                pct(count / metrics["post_exits"]),
            ]
        )
    add_table(
        doc,
        ["Time After Go Live", "Post Exits", "% of Post Exits"],
        bucket_rows,
        [Inches(2.7), Inches(1.5), Inches(1.6)],
        number_cols={1, 2},
    )
    add_paragraph(
        doc,
        "More than half of post-training exits happened within 14 days after Go Live. This makes the first two weeks after implementation the most important window for retention action.",
    )

    add_heading(doc, "Recommended Focus Areas", 1)
    add_bullet(doc, "For training, prioritize Wave 35 first, then Waves 33 and 34, because they concentrate most training exits.")
    add_bullet(doc, "For post-training, prioritize Wave 31 and Wave 34 by volume, and review Wave 27 because its post-training rate is high despite a small base.")
    add_bullet(doc, "Strengthen the first two weeks after Go Live with closer supervisor check-ins, expectation reinforcement, and fast issue resolution.")
    add_bullet(doc, "For voluntary exits, check job expectations, schedule fit, compensation perception, leadership support, and early employee experience.")
    add_bullet(doc, "For involuntary exits, review training pass criteria, early performance standards, nesting support, and quality coaching.")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
